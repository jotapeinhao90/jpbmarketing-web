from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def header_row(table, headers, bg='1F4E79'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_row(table, values, font_size=9):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        runs = cell.paragraphs[0].runs
        run = runs[0] if runs else cell.paragraphs[0].add_run('')
        run.font.size = Pt(font_size)
    return row

def h1(doc, text, color='1F4E79'):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor.from_string(color)
    return p

def h2(doc, text, color='2E75B6'):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(color)
    return p

def bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ══════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Plan de Crecimiento B2B — ETP Polimer')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string('1F4E79')
for line in ['Agencia: JPB Marketing SPA', 'Mayo 2026 — Hoja de ruta 6 meses',
             'Sitios: polymer.cl  ·  polimer-etp.cl (nuevo, split ads ~50%)']:
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(line).font.size = Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECCIÓN 1 — ESTADO GENERAL
# ══════════════════════════════════════════════════════════
h1(doc, '1. Estado General del Plan')

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
header_row(t, ['Fase', 'Descripción', 'Período', 'Estado'])
fases = [
    ('Fase 1', 'Diagnóstico y rediseño de campañas', 'Mes 1',   '🔄 En curso'),
    ('Fase 2', 'Infraestructura de captación B2B',   'Mes 2–3', '⏳ Pendiente'),
    ('Fase 3', 'Outbound y contenido sectorial',     'Mes 3–5', '⏳ Pendiente'),
    ('Fase 4', 'Optimización y escalamiento',        'Mes 5–6', '⏳ Pendiente'),
]
fase_colors = ['D9EAD3', 'FCE5CD', 'CFE2F3', 'E9D7FD']
for i, f in enumerate(fases):
    row = add_row(t, f)
    for cell in row.cells:
        set_cell_bg(cell, fase_colors[i])
doc.add_paragraph()

h2(doc, 'Tabla de tareas')
t2 = doc.add_table(rows=1, cols=5)
t2.style = 'Table Grid'
header_row(t2, ['#', 'Tarea', 'Fase', 'Complejidad', 'Estado'])
all_tasks = [
    (1,  'Auditoría profunda de campañas actuales',                      'F1', 'Media', '✅ Listo'),
    (2,  'Construir perfil ICP',                                         'F1', 'Alta',  '✅ Listo'),
    (3,  'Reestructurar campañas: match types + negative keywords',      'F1', 'Media', '✅ Listo'),
    (4,  'Customer Match + audiencias de intención personalizada',       'F1', 'Alta',  '⏳ Pendiente'),
    (5,  'Enhanced conversions + micro-conversiones',                    'F1', 'Alta',  '⏳ Pendiente'),
    (6,  'Landing page dedicada para empresas grandes (polimer-etp.cl)', 'F2', 'Alta',  '⏳ Pendiente'),
    (7,  'Lead magnet B2B: calculadora de ahorro en empaques',           'F2', 'Alta',  '⏳ Pendiente'),
    (8,  'Campaña Performance Max con señales B2B',                      'F2', 'Alta',  '⏳ Pendiente'),
    (9,  'Remarketing B2B por segmentos de intención',                   'F2', 'Media', '⏳ Pendiente'),
    (10, 'CRM básico (HubSpot Free o Pipedrive)',                        'F2', 'Media', '⏳ Pendiente'),
    (11, 'Secuencia email nurturing (3 correos)',                        'F2', 'Media', '⏳ Pendiente'),
    (12, 'Base de datos 200 empresas objetivo',                          'F3', 'Alta',  '⏳ Pendiente'),
    (13, 'Secuencia outreach LinkedIn para jefes de compras',            'F3', 'Alta',  '⏳ Pendiente'),
    (14, 'Contenido sectorial: artículos SEO',                           'F3', 'Media', '⏳ Pendiente'),
    (15, 'Campaña Display + YouTube por industria',                      'F3', 'Alta',  '⏳ Pendiente'),
    (16, 'Eventos industriales: Expomin, FIA, ferias packaging',         'F3', 'Media', '⏳ Pendiente'),
    (17, 'Alianzas con diseñadores de packaging e imprentas',            'F3', 'Alta',  '⏳ Pendiente'),
    (18, 'Búsquedas específicas por industria en Google Ads',            'F3', 'Media', '⏳ Pendiente'),
    (19, 'Análisis atribución multi-touch',                              'F4', 'Alta',  '⏳ Pendiente'),
    (20, 'Smart Bidding con valor diferenciado por tamaño empresa',      'F4', 'Alta',  '⏳ Pendiente'),
    (21, 'Caso de estudio cliente grande existente',                     'F4', 'Media', '⏳ Pendiente'),
    (22, 'Dashboard Looker Studio métricas comerciales',                 'F4', 'Media', '⏳ Pendiente'),
    (23, 'Prueba A/B landing pages (3 variantes)',                       'F4', 'Alta',  '⏳ Pendiente'),
    (24, 'Escalar canales ganadores',                                    'F4', 'Media', '⏳ Pendiente'),
    (25, 'Informe estratégico trimestral',                               'F4', 'Media', '⏳ Pendiente'),
]
fase_bg = {'F1': 'D9EAD3', 'F2': 'FCE5CD', 'F3': 'CFE2F3', 'F4': 'E9D7FD'}
for task in all_tasks:
    row = add_row(t2, task)
    set_cell_bg(row.cells[2], fase_bg[task[2]])
    comp_run = row.cells[3].paragraphs[0].runs[0]
    if task[3] == 'Alta':
        comp_run.font.color.rgb = RGBColor.from_string('C00000')
        comp_run.bold = True
    if '✅' in task[4]:
        row.cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string('375623')
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECCIÓN 2 — DIAGNÓSTICO
# ══════════════════════════════════════════════════════════
h1(doc, '2. Diagnóstico — Estado Actual de Campañas')

p = doc.add_paragraph()
p.add_run('Presupuesto mensual: ').bold = True
p.add_run('$500 USD  ·  ')
p.add_run('Contactos estimados: ').bold = True
p.add_run('~5 por día (~150/mes)  ·  ')
p.add_run('Perfil actual de leads: ').bold = True
p.add_run('Principalmente emprendimientos, con algunas empresas medianas/grandes')
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

h2(doc, 'Hallazgos críticos')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
header_row(t, ['Área', 'Hallazgo', 'Impacto'])
hallazgos = [
    ('Negative keywords', 'No existe ninguna lista de negativos', '🔴 Crítico — presupuesto desperdiciado en tráfico irrelevante'),
    ('Tracking',          'GA4 no instalado. Conversiones en Google Ads sin confirmar', '🔴 Crítico — se está optimizando a ciegas'),
    ('Formulario',        'Solo pide nombre y mail — no califica el lead', '🔴 Crítico — no se puede distinguir empresa grande de emprendedor'),
    ('Campañas',          'Sin separación entre búsqueda y Display/PMax', '🟠 Alto — no se sabe qué canal genera qué resultado'),
    ('Extensiones',       'Sin callouts ni sitelinks configurados', '🟡 Medio — se pierde espacio en el anuncio y CTR potencial'),
    ('Sitio web',         'La home no comunica claramente propuesta B2B/volumen', '🟡 Medio — empresas grandes no se identifican como el cliente objetivo'),
    ('Copies',            'Los anuncios sí mencionan volumen e industria', '✅ OK'),
    ('Velocidad sitio',   'Carga en menos de 3 segundos en móvil', '✅ OK'),
    ('Broad Match',       'No hay broad match sin restricción', '✅ OK'),
]
impacto_colors = {
    '🔴': 'FFC7CE',
    '🟠': 'FFDDB5',
    '🟡': 'FFEB9C',
    '✅': 'C6EFCE',
}
for h in hallazgos:
    row = add_row(t, h)
    color = next((v for k, v in impacto_colors.items() if k in h[2]), None)
    if color:
        set_cell_bg(row.cells[2], color)
doc.add_paragraph()

h2(doc, 'Acciones inmediatas (esta semana)')
acciones = [
    ('1 — URGENTE', 'Subir lista de negative keywords a Google Ads. Aplicar a todas las campañas. Estimado: recuperar 15–25% del presupuesto mensual.'),
    ('2 — URGENTE', 'Instalar GA4 en polymer.cl y configurar seguimiento de conversiones en Google Ads (formulario enviado = conversión).'),
    ('3 — Esta semana', 'Agregar campo al formulario: "¿Cuántas unidades necesitás por mes?" — esto califica automáticamente cada lead sin cambiar nada en las campañas.'),
    ('4 — Este mes', 'Separar campañas: crear campaña de búsqueda pura (keywords B2B) y dejar Display/PMax separado con su propio presupuesto.'),
    ('5 — Este mes', 'Agregar extensiones de anuncio: al menos 3 callouts ("Fabricación en Chile", "Pedidos desde X toneladas", "Entrega a todo el país") y 2 sitelinks.'),
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
header_row(t2, ['Prioridad', 'Acción'])
for accion in acciones:
    row = add_row(t2, accion)
    if 'URGENTE' in accion[0]:
        set_cell_bg(row.cells[0], 'FFC7CE')
        row.cells[0].paragraphs[0].runs[0].bold = True
    elif 'Esta semana' in accion[0]:
        set_cell_bg(row.cells[0], 'FFEB9C')
    else:
        set_cell_bg(row.cells[0], 'D9EAD3')
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECCIÓN 3 — ICP
# ══════════════════════════════════════════════════════════
h1(doc, '3. ICP — Perfil del Cliente Ideal')

h2(doc, 'Industrias objetivo')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
header_row(t, ['#', 'Industria', 'Por qué es ideal'])
for ind in [
    ('1', 'Retail / Supermercados',      'Alto volumen, compras recurrentes, estandarización de empaques'),
    ('2', 'Agroindustria / Exportación', 'Volúmenes masivos, requerimientos técnicos (resistencia, UV)'),
    ('3', 'Alimentos y bebidas',         'Normativa estricta = fidelidad al proveedor, alto consumo mensual'),
    ('4', 'Farmacia / Salud',            'Calidad certificada, bajo riesgo de cambio una vez dentro'),
    ('5', 'Construcción / Industria',    'Bolsas para cemento, áridos, insumos. Volumen estable'),
]:
    add_row(t, ind)
doc.add_paragraph()

h2(doc, 'Firmografía')
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
header_row(t, ['Criterio', 'Valor'])
for f in [
    ('Tamaño',            '+50 empleados (idealmente +200)'),
    ('Facturación anual', '+$500MM CLP'),
    ('Volumen de compra', '+500.000 unidades/mes o +5 toneladas/mes'),
    ('Tipo de compra',    'Recurrente, no puntual'),
    ('Proceso de compra', 'Formal, con área de abastecimiento definida'),
]:
    add_row(t, f)
doc.add_paragraph()

h2(doc, 'Perfil del decisor')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
header_row(t, ['Cargo', 'Rol', 'Motivación'])
for d in [
    ('Jefe de Compras / Abastecimiento', 'Decisor principal',    'Precio por tonelada, entrega, confiabilidad'),
    ('Gerente de Operaciones',           'Influenciador técnico', 'Calidad, estándares, normativa'),
    ('Gerente General / Financiero',     'Aprobador final',       'ROI, costo total, proveedor confiable'),
]:
    add_row(t, d)
doc.add_paragraph()

h2(doc, 'Calificación del lead')
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
header_row(t, ['Nivel', 'Señal', 'Acción'])
leads = [
    ('🔥 Caliente', 'Busca volumen/toneladas, visita precios >1 vez, descarga catálogo, tiempo >3 min, viene de retail/alimentos/agro/farmacia', 'Contactar en menos de 24h'),
    ('🌡 Tibio',    'Visita desde LinkedIn o Display, lee blog, abre emails de nurturing', 'Secuencia nurturing automática'),
    ('❄ Frío',     'Empresa <10 empleados, busca "bolsa regalo/ecológica/artesanal", pide precio unitario', 'Descartar / filtrar con negativo'),
]
lead_colors = ['FFE599', 'FCE5CD', 'CFE2F3']
for i, l in enumerate(leads):
    row = add_row(t, l)
    set_cell_bg(row.cells[0], lead_colors[i])
doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECCIÓN 4 — NEGATIVE KEYWORDS
# ══════════════════════════════════════════════════════════
h1(doc, '4. Negative Keywords — Google Ads')
p = doc.add_paragraph()
p.add_run('Subir como lista compartida: ').bold = True
p.add_run('Google Ads → Herramientas → Biblioteca compartida → Listas de palabras clave negativas → Crear lista "Polimer Filtro B2B v1" → Aplicar a todas las campañas.')
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

bloques = [
    ('Tipo de producto no deseado',
     ['bolsa tela', 'bolsa ecológica', 'bolsa biodegradable', 'bolsa papel', 'bolsa kraft',
      'bolsa algodón', 'bolsa yute', 'bolsa reutilizable', 'bolsita', 'bolsas pequeñas',
      'bolsas de regalo', 'bolsa para regalo', 'bolsas personalizadas pequeñas',
      'bolsas boda', 'bolsas cumpleaños', 'bolsas fiesta', 'bolsas navideña']),
    ('Uso doméstico / particular',
     ['bolsa basura', 'bolsa de basura', 'bolsa aspiradora', 'bolsa congelador', 'bolsa freezer',
      'bolsa zip', 'bolsa para ropa', 'bolsa de viaje', 'bolsa de playa', 'bolsas para mascotas']),
    ('Volumen bajo / retail unitario',
     ['precio unitario', 'precio por unidad', 'precio una bolsa', 'comprar una bolsa',
      'sin mínimo', 'pedir 100 bolsas', '50 unidades', 'pedido pequeño', 'muestra gratis']),
    ('Artesanal / DIY / emprendimiento',
     ['artesanal', 'hecho a mano', 'casero', 'DIY', 'emprendimiento',
      'negocio pequeño', 'microempresa', 'feria artesanal']),
    ('Segunda mano / clasificados',
     ['usado', 'segunda mano', 'mercado libre bolsas', 'yapo bolsas', 'olx bolx']),
    ('Informativos sin intención de compra',
     ['que es polietileno', 'historia del polietileno', 'como reciclar bolsas',
      'contaminación plástica', 'impacto ambiental bolsas', 'alternativas al plástico']),
]
for titulo, keywords in bloques:
    h2(doc, titulo)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    header_row(t, ['Keyword negativa', 'Concordancia'])
    for kw in keywords:
        add_row(t, [kw, 'Phrase match'])
    doc.add_paragraph()

doc.save('/Users/jotape/Gravity/polimer-etp/Plan_Polimer_B2B.docx')
print("✅ Plan_Polimer_B2B.docx actualizado")
